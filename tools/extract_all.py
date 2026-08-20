import sys
import io
import os
import re
import glob

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

BASE = r"D:\AI\projects\mtms-aqua-haier-kb\materi-drive"
OUT = r"D:\AI\projects\mtms-aqua-haier-kb\tools\extracted"
os.makedirs(OUT, exist_ok=True)

def clean(t):
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()

def extract_pdf(path, outname):
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            t = page.extract_text() or ""
            parts.append(f"--- PAGE {i} ---\n{t}")
    txt = clean("\n".join(parts))
    if len(txt) > 50:
        open(os.path.join(OUT, outname), "w", encoding="utf-8").write(txt)
        return len(txt)
    return 0

def extract_docx(path, outname):
    import docx
    d = docx.Document(path)
    parts = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for t_idx, table in enumerate(d.tables, 1):
        parts.append(f"--- TABLE {t_idx} ---")
        for row in table.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            parts.append(" | ".join(cells))
    txt = clean("\n".join(parts))
    if len(txt) > 50:
        open(os.path.join(OUT, outname), "w", encoding="utf-8").write(txt)
        return len(txt)
    return 0

def extract_xlsx(path, outname):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"=== SHEET: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            vals = ["" if v is None else str(v).replace("\n", " ") for v in row]
            line = " | ".join(vals).rstrip(" |")
            if line.strip():
                parts.append(line)
    txt = clean("\n".join(parts))
    if len(txt) > 50:
        open(os.path.join(OUT, outname), "w", encoding="utf-8").write(txt)
        return len(txt)
    return 0

def extract_xlsb(path, outname):
    from pyxlsb import open_workbook
    parts = []
    with open_workbook(path) as wb:
        for sheet in wb.sheets:
            parts.append(f"=== SHEET: {sheet} ===")
            with wb.get_sheet(sheet) as ws:
                for row in ws.rows():
                    vals = ["" if v is None else str(v) for v in row]
                    line = " | ".join(vals).rstrip(" |")
                    if line.strip():
                        parts.append(line)
    txt = clean("\n".join(parts))
    if len(txt) > 50:
        open(os.path.join(OUT, outname), "w", encoding="utf-8").write(txt)
        return len(txt)
    return 0

def main():
    results = []
    errors = []
    for root, dirs, fns in os.walk(BASE):
        for fn in fns:
            src = os.path.join(root, fn)
            lower = fn.lower()
            rel = os.path.relpath(src, BASE).replace(os.sep, "__").replace(" ", "_")
            try:
                if lower.endswith(".pdf"):
                    n = extract_pdf(src, rel + ".txt")
                elif lower.endswith(".docx"):
                    n = extract_docx(src, rel + ".txt")
                elif lower.endswith(".xlsx"):
                    n = extract_xlsx(src, rel + ".txt")
                elif lower.endswith(".xlsb"):
                    n = extract_xlsb(src, rel + ".txt")
                else:
                    continue
                results.append((n, rel))
            except Exception as e:
                errors.append((rel, str(e)[:150]))
    for n, rel in sorted(results, reverse=True):
        print(f"{n:>10}  {rel}")
    print(f"\nTOTAL {len(results)} extracted, empty: {sum(1 for n,_ in results if n==0)}")
    if errors:
        print(f"\nERRORS ({len(errors)}):")
        for rel, e in errors:
            print(f"  ERR {rel}: {e}")

if __name__ == "__main__":
    main()